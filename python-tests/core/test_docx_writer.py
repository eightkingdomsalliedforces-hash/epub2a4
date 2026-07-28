from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from epub_a4_word.docx_writer import write_docx
from epub_a4_word.models import ImageBlock, TextBlock, TextRun
from epub_a4_word.pagination import LayoutSettings, MiniPage


def _page(text: str) -> MiniPage:
    return MiniPage([TextBlock((TextRun(text),), style="body")], used_points=20)


def test_writer_creates_one_two_by_two_table_per_four_mini_pages(tmp_path: Path) -> None:
    output = tmp_path / "four-up.docx"
    pages = [_page(f"頁 {index}") for index in range(1, 6)]

    write_docx(pages, output, resources={}, media_types={}, settings=LayoutSettings())

    document = Document(output)
    assert len(document.tables) == 2
    assert all(len(table.rows) == 2 for table in document.tables)
    assert all(len(row.cells) == 2 for table in document.tables for row in table.rows)
    text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "頁 1" in text
    assert "頁 5" in text


def test_writer_embeds_image_media(tmp_path: Path) -> None:
    image_data = BytesIO()
    Image.new("RGB", (200, 300), "white").save(image_data, format="PNG")
    output = tmp_path / "image.docx"
    image = ImageBlock("Images/test.png", alt_text="測試圖片")

    write_docx(
        [MiniPage([image], used_points=100)],
        output,
        resources={"Images/test.png": image_data.getvalue()},
        media_types={"Images/test.png": "image/png"},
        settings=LayoutSettings(),
    )

    with ZipFile(output) as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
    assert len(media) == 1


def test_writer_renders_consecutive_tables_without_blank_pages(tmp_path: Path) -> None:
    import shutil
    import subprocess

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    pdfinfo = shutil.which("pdfinfo")
    if not soffice or not pdfinfo:
        import pytest
        pytest.skip("LibreOffice/pdfinfo not installed")

    output = tmp_path / "two-sheets.docx"
    pages = [_page(f"頁 {index}") for index in range(1, 6)]
    write_docx(pages, output, resources={}, media_types={}, settings=LayoutSettings())
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(output)],
        check=True,
        capture_output=True,
        text=True,
    )
    info = subprocess.run([pdfinfo, str(tmp_path / "two-sheets.pdf")], check=True, capture_output=True, text=True).stdout
    page_line = next(line for line in info.splitlines() if line.startswith("Pages:"))
    assert int(page_line.split(":", 1)[1].strip()) == 2


def test_writer_uses_signature16_cell_order(tmp_path: Path) -> None:
    output = tmp_path / "signature.docx"
    pages = [_page(f"P{index}") for index in range(1, 17)]

    write_docx(
        pages,
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(),
        imposition_mode="signature16",
    )

    document = Document(output)
    assert len(document.tables) == 4
    orders = [
        [cell.paragraphs[0].text for row in table.rows for cell in row.cells]
        for table in document.tables
    ]
    assert orders == [
        ["P16", "P1", "P14", "P3"],
        ["P2", "P15", "P4", "P13"],
        ["P12", "P5", "P10", "P7"],
        ["P6", "P11", "P8", "P9"],
    ]


def test_writer_renders_signature16_as_four_a4_sides(tmp_path: Path) -> None:
    import shutil
    import subprocess

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    pdfinfo = shutil.which("pdfinfo")
    if not soffice or not pdfinfo:
        import pytest
        pytest.skip("LibreOffice/pdfinfo not installed")

    output = tmp_path / "signature.docx"
    pages = [_page(f"P{index}") for index in range(1, 17)]
    write_docx(
        pages,
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(),
        imposition_mode="signature16",
    )
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(output)],
        check=True,
        capture_output=True,
        text=True,
    )
    info = subprocess.run([pdfinfo, str(tmp_path / "signature.pdf")], check=True, capture_output=True, text=True).stdout
    page_line = next(line for line in info.splitlines() if line.startswith("Pages:"))
    assert int(page_line.split(":", 1)[1].strip()) == 4


def test_writer_displays_logical_number_on_text_and_image_pages(tmp_path: Path) -> None:
    output = tmp_path / "logical-numbers.docx"
    text_page = MiniPage(
        [TextBlock((TextRun("序章正文"),), style="body")],
        used_points=20,
        logical_page_number=1,
    )
    image_page = MiniPage(
        [ImageBlock("Images/test.png")],
        used_points=20,
        logical_page_number=2,
    )
    image_data = BytesIO()
    Image.new("RGB", (200, 300), "white").save(image_data, format="PNG")

    write_docx(
        [text_page, image_page],
        output,
        resources={"Images/test.png": image_data.getvalue()},
        media_types={"Images/test.png": "image/png"},
        settings=LayoutSettings(page_numbers=True, imposition_mode="four_up"),
    )

    document = Document(output)
    cells = [cell for row in document.tables[0].rows for cell in row.cells]
    assert "1" in cells[0].text.splitlines()
    assert "2" in cells[1].text.splitlines()


def test_writer_page_number_switch_removes_numbers_from_all_page_types(
    tmp_path: Path,
) -> None:
    output = tmp_path / "no-logical-numbers.docx"
    pages = [
        MiniPage(
            [TextBlock((TextRun("正文"),), style="body")],
            logical_page_number=1,
        ),
        MiniPage([ImageBlock("Images/test.png")], logical_page_number=2),
    ]
    image_data = BytesIO()
    Image.new("RGB", (200, 300), "white").save(image_data, format="PNG")

    write_docx(
        pages,
        output,
        resources={"Images/test.png": image_data.getvalue()},
        media_types={"Images/test.png": "image/png"},
        settings=LayoutSettings(page_numbers=False, imposition_mode="four_up"),
    )

    cells = [
        cell
        for row in Document(output).tables[0].rows
        for cell in row.cells
    ]
    assert "1" not in cells[0].text.splitlines()
    assert "2" not in cells[1].text.splitlines()


def test_right_binding_mirrors_page_number_alignment(tmp_path: Path) -> None:
    output = tmp_path / "right-binding-page-numbers.docx"
    pages = [
        MiniPage(
            [TextBlock((TextRun("第一頁"),), style="body")],
            logical_page_number=1,
        ),
        MiniPage(
            [TextBlock((TextRun("第二頁"),), style="body")],
            logical_page_number=2,
        ),
    ]

    write_docx(
        pages,
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(
            imposition_mode="single_a5",
            binding_direction="right",
            page_numbers=True,
        ),
        imposition_mode="single_a5",
    )

    rows = Document(output).tables[0].rows
    assert rows[0].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert rows[1].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


def test_borderless_writer_sets_zero_a4_outer_margins(tmp_path: Path) -> None:
    output = tmp_path / "borderless.docx"
    page = MiniPage(
        [TextBlock((TextRun("序章正文"),), style="body")],
        used_points=20,
        logical_page_number=1,
    )

    write_docx(
        [page],
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(margin_mode="borderless", imposition_mode="four_up"),
    )

    document = Document(output)
    section = document.sections[0]
    assert section.top_margin.cm == 0
    assert section.bottom_margin.cm == 0
    assert section.left_margin.cm == 0
    assert section.right_margin.cm == 0


def test_writer_renders_page_number_even_when_text_page_is_dense(tmp_path: Path) -> None:
    import shutil
    import subprocess

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    pdftotext = shutil.which("pdftotext")
    if not soffice or not pdftotext:
        import pytest
        pytest.skip("LibreOffice/pdftotext not installed")

    output = tmp_path / "dense-number.docx"
    dense = TextBlock((TextRun("這是正文。" * 900),), style="body")
    page = MiniPage([dense], used_points=300, logical_page_number=9876)
    write_docx(
        [page],
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(
            margin_mode="borderless",
            imposition_mode="four_up",
            page_numbers=True,
        ),
    )

    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(output)],
        check=True,
        capture_output=True,
        text=True,
    )
    extracted = subprocess.run(
        [pdftotext, str(tmp_path / "dense-number.pdf"), "-"],
        check=True,
        capture_output=True,
        text=True,
    ).stdout
    assert "9876" in extracted


def test_writer_creates_one_a5_page_per_content_page(tmp_path: Path) -> None:
    output = tmp_path / "a5.docx"
    pages = [_page("A5 第一頁"), _page("A5 第二頁")]

    write_docx(
        pages,
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(imposition_mode="single_a5"),
        imposition_mode="single_a5",
    )

    document = Document(output)
    section = document.sections[0]
    assert abs(section.page_width.cm - 14.8) < 0.02
    assert abs(section.page_height.cm - 21.0) < 0.02
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 2
    assert all(len(row.cells) == 1 for row in table.rows)
    assert [row.cells[0].text.strip() for row in table.rows] == ["A5 第一頁", "A5 第二頁"]


def test_writer_creates_one_4x6_page_per_content_page(tmp_path: Path) -> None:
    output = tmp_path / "4x6.docx"
    pages = [_page("4×6 第一頁"), _page("4×6 第二頁")]

    write_docx(
        pages,
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(imposition_mode="single_4x6"),
        imposition_mode="single_4x6",
    )

    document = Document(output)
    section = document.sections[0]
    assert abs(section.page_width.cm - 10.16) < 0.02
    assert abs(section.page_height.cm - 15.24) < 0.02
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 2
    assert all(len(row.cells) == 1 for row in table.rows)
    assert [row.cells[0].text.strip() for row in table.rows] == ["4×6 第一頁", "4×6 第二頁"]


def test_writer_renders_single_a5_without_blank_pages(tmp_path: Path) -> None:
    import shutil
    import subprocess

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    pdfinfo = shutil.which("pdfinfo")
    if not soffice or not pdfinfo:
        import pytest
        pytest.skip("LibreOffice/pdfinfo not installed")

    output = tmp_path / "a5-render.docx"
    pages = [_page(f"A5 P{index}") for index in range(1, 4)]
    write_docx(
        pages,
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(imposition_mode="single_a5"),
        imposition_mode="single_a5",
    )
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(output)],
        check=True,
        capture_output=True,
        text=True,
    )
    info = subprocess.run([pdfinfo, str(tmp_path / "a5-render.pdf")], check=True, capture_output=True, text=True).stdout
    page_line = next(line for line in info.splitlines() if line.startswith("Pages:"))
    assert int(page_line.split(":", 1)[1].strip()) == 3


def test_writer_uses_exact_shared_line_height_for_body_and_heading(
    tmp_path: Path,
) -> None:
    from lxml import etree

    output = tmp_path / "exact-line-height.docx"
    page = MiniPage(
        [
            TextBlock((TextRun("正文 A Certain Magical Index"),), style="body"),
            TextBlock((TextRun("標題"),), style="heading"),
        ],
        used_points=40,
    )

    write_docx(
        [page],
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(page_numbers=False, imposition_mode="single_a5"),
        imposition_mode="single_a5",
    )

    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(output) as archive:
        root = etree.fromstring(archive.read("word/document.xml"))
    paragraphs = root.xpath(".//w:tc/w:p", namespaces=namespace)
    assert len(paragraphs) == 2

    body_spacing = paragraphs[0].find("w:pPr/w:spacing", namespaces=namespace)
    heading_spacing = paragraphs[1].find("w:pPr/w:spacing", namespaces=namespace)
    assert body_spacing is not None
    assert heading_spacing is not None
    word_ns = namespace["w"]
    assert body_spacing.get(f"{{{word_ns}}}lineRule") == "exact"
    assert int(body_spacing.get(f"{{{word_ns}}}line")) == 230
    assert int(body_spacing.get(f"{{{word_ns}}}after")) == 50
    assert heading_spacing.get(f"{{{word_ns}}}lineRule") == "exact"
    assert int(heading_spacing.get(f"{{{word_ns}}}line")) == 290
    assert int(heading_spacing.get(f"{{{word_ns}}}after")) == 100


def test_pre_paginated_heading_does_not_ask_word_to_keep_next(tmp_path: Path) -> None:
    output = tmp_path / "heading-no-keep-next.docx"
    page = MiniPage(
        [
            TextBlock((TextRun("第一章"),), style="heading"),
            TextBlock((TextRun("正文內容"),), style="body"),
        ],
        used_points=40,
    )

    write_docx(
        [page],
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(imposition_mode="single_a5"),
        imposition_mode="single_a5",
    )

    with ZipFile(output) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")

    assert "<w:keepNext" not in document_xml


def test_vertical_writer_emits_native_tb_rl_and_keeps_source_text(
    tmp_path: Path,
) -> None:
    output = tmp_path / "vertical.docx"
    text = "中文 English 2026"

    write_docx(
        [
            MiniPage(
                [TextBlock((TextRun(text),), style="body")],
                logical_page_number=1,
            )
        ],
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(
            writing_mode="taiwan_vertical",
            binding_direction="right",
            page_numbers=False,
        ),
    )

    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml")
    assert b'<w:textDirection w:val="tbRl"' in xml
    rendered_text = "\n".join(
        cell.text
        for row in Document(output).tables[0].rows
        for cell in row.cells
    )
    assert text in rendered_text


def test_horizontal_writer_does_not_emit_vertical_text_direction(
    tmp_path: Path,
) -> None:
    output = tmp_path / "horizontal.docx"

    write_docx(
        [_page("中文 English 2026")],
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(
            writing_mode="horizontal",
            binding_direction="left",
        ),
    )

    with ZipFile(output) as archive:
        assert b'<w:textDirection w:val="tbRl"' not in archive.read(
            "word/document.xml"
        )


def test_vertical_writer_keeps_image_in_horizontal_nested_cell(
    tmp_path: Path,
) -> None:
    image_data = BytesIO()
    Image.new("RGB", (200, 300), "white").save(image_data, format="PNG")
    output = tmp_path / "vertical-image.docx"

    write_docx(
        [MiniPage([ImageBlock("Images/test.png")], logical_page_number=1)],
        output,
        resources={"Images/test.png": image_data.getvalue()},
        media_types={"Images/test.png": "image/png"},
        settings=LayoutSettings(
            writing_mode="taiwan_vertical",
            binding_direction="right",
        ),
    )

    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml")
        names = archive.namelist()
    assert b'<w:textDirection w:val="tbRl"' in xml
    assert b'<w:textDirection w:val="lrTb"' in xml
    assert len(Document(output).inline_shapes) == 1
    assert any(name.startswith("word/media/") for name in names)


def test_vertical_writer_returns_word_compatibility_warning(
    tmp_path: Path,
) -> None:
    warnings = write_docx(
        [_page("直排")],
        tmp_path / "warning.docx",
        resources={},
        media_types={},
        settings=LayoutSettings(writing_mode="taiwan_vertical"),
    )

    assert any(
        "Microsoft Word" in warning and "East Asia" in warning
        for warning in warnings
    )
