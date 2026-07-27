from __future__ import annotations

from dataclasses import replace
from io import BytesIO
from pathlib import Path
from typing import Mapping, Sequence

from PIL import Image, UnidentifiedImageError
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .crop_marks import add_page_guides
from .imposition import ImpositionMode, build_imposition
from .models import ImageBlock, TextBlock, TextRun
from .page_placement import build_page_placement
from .pagination import LayoutSettings, MiniPage, resolve_layout
from .text_metrics import paragraph_metrics

TWIPS_PER_CM = 1440.0 / 2.54
_SINGLE_PAGE_TABLE_MODES = frozenset({"single_a5", "single_4x6", "b6_on_a5"})


def _cm_to_twips(value: float) -> int:
    return max(0, round(value * TWIPS_PER_CM))


def _set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, visible: bool) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        is_center_guide = edge in {"insideH", "insideV"}
        element.set(qn("w:val"), "dashed" if visible and is_center_guide else "nil")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "909090")


def _set_fixed_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_table_width(table, width_cm: float) -> None:
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), str(_cm_to_twips(width_cm)))
    width.set(qn("w:type"), "dxa")


def _set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)


def _clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def _configure_page_prefix(paragraph, page_break_before: bool) -> None:
    _clear_paragraph(paragraph)
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = Pt(1)
    fmt.page_break_before = page_break_before
    run = paragraph.add_run("\u200b")
    run.font.size = Pt(1)


def _normalized_image(data: bytes, media_type: str, resource_path: str) -> tuple[BytesIO, int, int]:
    suffix = Path(resource_path).suffix.lower()
    if media_type == "image/svg+xml" or suffix == ".svg":
        try:
            import cairosvg
        except (ImportError, OSError) as exc:
            raise ValueError("SVG 圖片需要額外安裝 CairoSVG 與其系統元件") from exc
        stream = BytesIO(cairosvg.svg2png(bytestring=data))
    else:
        stream = BytesIO(data)
    try:
        with Image.open(stream) as image:
            image.load()
            width_px, height_px = image.size
            fmt = (image.format or "").upper()
            if fmt not in {"JPEG", "PNG", "GIF", "TIFF", "BMP"}:
                converted = BytesIO()
                if image.mode not in {"RGB", "RGBA"}:
                    image = image.convert("RGBA")
                image.save(converted, format="PNG")
                stream = converted
            else:
                stream.seek(0)
            return stream, width_px, height_px
    except UnidentifiedImageError as exc:
        raise ValueError(f"無法讀取圖片：{resource_path}") from exc


def _add_text_block(cell, block: TextBlock, settings: LayoutSettings, first: bool) -> None:
    paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
    if first:
        _clear_paragraph(paragraph)
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    font_size = settings.heading_font_pt if block.style == "heading" else settings.body_font_pt
    metrics = paragraph_metrics(
        font_size,
        settings.line_spacing,
        settings.heading_spacing_pt
        if block.style == "heading"
        else settings.paragraph_spacing_pt,
    )
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(metrics.line_height_pt)
    fmt.space_after = Pt(metrics.spacing_after_pt)
    if block.style == "heading":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif block.style == "quote":
        fmt.left_indent = Pt(font_size * 1.5)
        fmt.right_indent = Pt(font_size * 1.5)
    else:
        fmt.first_line_indent = Pt(font_size * settings.first_line_indent_chars)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for source_run in block.runs:
        run = paragraph.add_run(source_run.text)
        run.bold = source_run.bold or block.style == "heading"
        run.italic = source_run.italic
        run.font.size = Pt(font_size)
        _set_east_asia_font(run, settings.font_name)


def _add_image_block(
    cell,
    block: ImageBlock,
    settings: LayoutSettings,
    resources: Mapping[str, bytes],
    media_types: Mapping[str, str],
    first: bool,
) -> str | None:
    data = resources.get(block.resource_path)
    if data is None:
        return f"缺少圖片資料：{block.resource_path}"
    try:
        stream, width_px, height_px = _normalized_image(
            data, media_types.get(block.resource_path, ""), block.resource_path
        )
    except Exception as exc:
        return str(exc)
    assert settings.max_image_width_pt is not None and settings.max_image_height_pt is not None
    width_pt = settings.max_image_width_pt
    height_pt = width_pt * height_px / max(1, width_px)
    if height_pt > settings.max_image_height_pt:
        height_pt = settings.max_image_height_pt
        width_pt = height_pt * width_px / max(1, height_px)
    paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
    if first:
        _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(settings.paragraph_spacing_pt)
    paragraph.add_run().add_picture(stream, width=Pt(width_pt), height=Pt(height_pt))
    return None


def _add_page_number(cell, page_number: int, settings: LayoutSettings, first: bool) -> None:
    paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
    if first:
        _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if page_number % 2 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = Pt(8)
    run = paragraph.add_run(str(page_number))
    run.font.size = Pt(6.5)
    _set_east_asia_font(run, settings.font_name)


def _populate_cell(
    cell,
    page: MiniPage | None,
    slot: int,
    grid_cols: int,
    settings: LayoutSettings,
    resources: Mapping[str, bytes],
    media_types: Mapping[str, str],
) -> list[str]:
    warnings: list[str] = []
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    assert settings.cell_outer_margin_cm is not None
    assert settings.gutter_margin_cm is not None
    assert settings.cell_vertical_margin_cm is not None
    if grid_cols == 1:
        start_cm = end_cm = settings.cell_outer_margin_cm
    else:
        column = slot % grid_cols
        start_cm = settings.cell_outer_margin_cm if column == 0 else settings.gutter_margin_cm
        end_cm = settings.gutter_margin_cm if column == 0 else settings.cell_outer_margin_cm
    vertical_twips = _cm_to_twips(settings.cell_vertical_margin_cm)
    _set_cell_margins(cell, vertical_twips, _cm_to_twips(start_cm), vertical_twips, _cm_to_twips(end_cm))
    if page is None:
        _clear_paragraph(cell.paragraphs[0])
        return warnings
    first = True
    if settings.page_numbers and page.logical_page_number is not None and page.has_text:
        _add_page_number(cell, page.logical_page_number, settings, first)
        first = False
    for block in page.blocks:
        if isinstance(block, TextBlock):
            _add_text_block(cell, block, settings, first)
            first = False
        elif isinstance(block, ImageBlock):
            warning = _add_image_block(cell, block, settings, resources, media_types, first)
            if warning:
                warnings.append(warning)
                _add_text_block(
                    cell,
                    TextBlock((TextRun(f"[圖片無法顯示：{block.alt_text or block.resource_path}]"),), style="body"),
                    settings,
                    first,
                )
            first = False
    return warnings


def write_docx(
    pages: Sequence[MiniPage],
    output_path: Path | str,
    *,
    resources: Mapping[str, bytes],
    media_types: Mapping[str, str],
    settings: LayoutSettings,
    title: str = "",
    author: str = "",
    imposition_mode: ImpositionMode = "four_up",
) -> list[str]:
    settings = resolve_layout(replace(settings, imposition_mode=imposition_mode))
    assert settings.cell_width_cm is not None and settings.cell_height_cm is not None
    assert settings.paper_width_cm is not None and settings.paper_height_cm is not None
    assert settings.grid_rows is not None and settings.grid_cols is not None
    assert settings.page_margin_left_cm is not None and settings.page_margin_right_cm is not None
    assert settings.page_margin_top_cm is not None and settings.page_margin_bottom_cm is not None

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(settings.paper_width_cm)
    section.page_height = Cm(settings.paper_height_cm)
    section.top_margin = Cm(settings.page_margin_top_cm)
    section.bottom_margin = Cm(settings.page_margin_bottom_cm)
    section.left_margin = Cm(settings.page_margin_left_cm)
    section.right_margin = Cm(settings.page_margin_right_cm)
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)
    document.core_properties.title = title
    document.core_properties.author = author

    placement = build_page_placement(settings)
    add_page_guides(
        section,
        placement.guides,
        paper_width_mm=placement.paper_width_mm,
        paper_height_mm=placement.paper_height_mm,
        render_mode=settings.guide_render_mode,
    )

    warnings: list[str] = []
    plan = build_imposition(len(pages), imposition_mode)

    def configure_table(table) -> None:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_fixed_layout(table)
        _set_table_width(table, settings.cell_width_cm * settings.grid_cols)
        _set_table_borders(table, False)

    def configure_row(row) -> None:
        row.height = Cm(settings.cell_height_cm)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell in row.cells:
            cell.width = Cm(settings.cell_width_cm)

    if imposition_mode in _SINGLE_PAGE_TABLE_MODES:
        if plan.sides:
            table = document.add_table(rows=len(plan.sides), cols=1)
            configure_table(table)
            for side_index, slots in enumerate(plan.sides):
                row = table.rows[side_index]
                configure_row(row)
                physical_page_number = slots[0]
                page = (
                    pages[physical_page_number - 1]
                    if physical_page_number is not None
                    else None
                )
                warnings.extend(
                    _populate_cell(
                        row.cells[0],
                        page,
                        0,
                        1,
                        settings,
                        resources,
                        media_types,
                    )
                )
    else:
        first_prefix = document.add_paragraph()
        for side_index, slots in enumerate(plan.sides):
            prefix = first_prefix if side_index == 0 else document.add_paragraph()
            _configure_page_prefix(prefix, page_break_before=side_index > 0)
            table = document.add_table(rows=settings.grid_rows, cols=settings.grid_cols)
            configure_table(table)
            for row in table.rows:
                configure_row(row)
            for slot, physical_page_number in enumerate(slots):
                cell = table.cell(slot // settings.grid_cols, slot % settings.grid_cols)
                page = (
                    pages[physical_page_number - 1]
                    if physical_page_number is not None
                    else None
                )
                warnings.extend(
                    _populate_cell(
                        cell,
                        page,
                        slot,
                        settings.grid_cols,
                        settings,
                        resources,
                        media_types,
                    )
                )
    document.save(output)
    return warnings
