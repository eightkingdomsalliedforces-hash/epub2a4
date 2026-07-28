from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Cm, Pt

from epub_a4_word.pagination import LayoutSettings
from epub_a4_word.word_reflow import convert_docx


def _make_source_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    section.header.paragraphs[0].text = "舊 A4 頁首"
    section.footer.paragraphs[0].text = "舊 A4 頁尾 99"

    first = document.add_paragraph()
    first.paragraph_format.first_line_indent = Cm(0.74)
    first.paragraph_format.space_after = Pt(8)
    first.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    first.add_run("第一句不應被拆成段落。第二句仍然在同一個 Word 段落。")

    second = document.add_paragraph()
    second.add_run("這是第二段，包含 ")
    bold = second.add_run("粗體")
    bold.bold = True
    second.add_run("、")
    italic = second.add_run("斜體")
    italic.italic = True
    second.add_run("與底線")
    second.runs[-1].underline = True
    second.add_run("。")

    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    table.cell(0, 0).text = "表格 A"
    table.cell(0, 1).text = "表格 B"
    table.cell(1, 0).text = "表格 C"
    table.cell(1, 1).text = "表格 D"
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(7.0)

    image = BytesIO()
    Image.new("RGB", (120, 80), "white").save(image, format="PNG")
    document.add_paragraph().add_run().add_picture(BytesIO(image.getvalue()), width=Cm(12.0))

    breaker = document.add_paragraph("手動換頁前")
    breaker.add_run().add_break(WD_BREAK.PAGE)
    document.add_paragraph("手動換頁後")
    document.save(path)


def _body_paragraph_texts(document: Document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs]


def _page_break_count(path: Path) -> int:
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    return xml.count(b'w:type="page"')


def test_convert_docx_reflows_to_a5_without_changing_paragraph_boundaries(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "a5.docx"
    _make_source_docx(source)

    source_document = Document(source)
    source_texts = _body_paragraph_texts(source_document)
    source_breaks = _page_break_count(source)

    result = convert_docx(
        source,
        output,
        LayoutSettings(imposition_mode="single_a5", margin_mode="safe", page_numbers=True),
    )

    converted = Document(output)
    assert _body_paragraph_texts(converted) == source_texts
    assert len(converted.paragraphs) == len(source_document.paragraphs)
    assert "第一句不應被拆成段落。第二句仍然在同一個 Word 段落。" in converted.paragraphs[0].text
    assert converted.paragraphs[0].paragraph_format.first_line_indent.cm == source_document.paragraphs[0].paragraph_format.first_line_indent.cm
    assert converted.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert converted.paragraphs[1].runs[1].bold is True
    assert converted.paragraphs[1].runs[3].italic is True
    assert converted.paragraphs[1].runs[4].underline is True
    assert len(converted.tables) == 1
    assert converted.tables[0].cell(1, 1).text == "表格 D"
    assert len(converted.inline_shapes) == 1
    assert _page_break_count(output) == source_breaks

    section = converted.sections[0]
    assert abs(section.page_width.cm - 14.8) < 0.02
    assert abs(section.page_height.cm - 21.0) < 0.02
    assert abs(section.top_margin.cm - 0.5) < 0.02
    assert abs(section.left_margin.cm - 0.5) < 0.02
    assert "舊 A4 頁首" not in "\n".join(p.text for p in section.header.paragraphs)
    assert "舊 A4 頁尾" not in "\n".join(p.text for p in section.footer.paragraphs)
    assert "PAGE" in section.footer._element.xml
    assert result.source_format == "docx"
    assert result.imposition_mode == "single_a5"


def test_convert_docx_reflows_to_exact_4x6_and_can_disable_new_page_numbers(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "photo.docx"
    _make_source_docx(source)

    convert_docx(
        source,
        output,
        LayoutSettings(imposition_mode="single_4x6", margin_mode="maximized", page_numbers=False),
    )

    converted = Document(output)
    section = converted.sections[0]
    assert abs(section.page_width.cm - 10.16) < 0.02
    assert abs(section.page_height.cm - 15.24) < 0.02
    assert abs(section.top_margin.cm - 0.2) < 0.02
    assert abs(section.bottom_margin.cm - 0.2) < 0.02
    assert "PAGE" not in section.footer._element.xml
    assert "舊 A4 頁尾" not in section.footer._element.xml
    available_width_cm = section.page_width.cm - section.left_margin.cm - section.right_margin.cm
    assert all(shape.width.cm <= available_width_cm + 0.02 for shape in converted.inline_shapes)
    grid_widths = [int(col.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w")) for col in converted.tables[0]._tbl.tblGrid.gridCol_lst]
    available_twips = round(available_width_cm * 1440 / 2.54)
    assert sum(grid_widths) <= available_twips + 2


def test_convert_docx_rejects_imposition_modes_that_are_not_single_page(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    _make_source_docx(source)

    import pytest

    with pytest.raises(ValueError, match="DOCX.*A5.*4×6"):
        convert_docx(
            source,
            tmp_path / "bad.docx",
            LayoutSettings(imposition_mode="signature16"),
        )


def test_convert_docx_vertical_sets_tb_rl_without_losing_content(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "vertical.docx"
    _make_source_docx(source)
    before = Document(source)

    result = convert_docx(
        source,
        output,
        LayoutSettings(
            imposition_mode="single_a5",
            writing_mode="taiwan_vertical",
            binding_direction="right",
            page_numbers=True,
        ),
    )

    converted = Document(output)
    assert _body_paragraph_texts(converted) == _body_paragraph_texts(before)
    assert converted.paragraphs[1].runs[1].bold is True
    assert converted.paragraphs[1].runs[3].italic is True
    assert len(converted.inline_shapes) == 1
    assert _page_break_count(output) == _page_break_count(source)
    with ZipFile(output) as archive:
        assert b'<w:textDirection w:val="tbRl"' in archive.read(
            "word/document.xml"
        )
    assert any("Microsoft Word" in warning for warning in result.warnings)


def test_convert_docx_horizontal_removes_tb_rl_from_sections(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    vertical = tmp_path / "vertical.docx"
    horizontal = tmp_path / "horizontal.docx"
    _make_source_docx(source)
    convert_docx(
        source,
        vertical,
        LayoutSettings(
            imposition_mode="single_a5",
            writing_mode="taiwan_vertical",
        ),
    )

    convert_docx(
        vertical,
        horizontal,
        LayoutSettings(
            imposition_mode="single_a5",
            writing_mode="horizontal",
            binding_direction="left",
        ),
    )

    with ZipFile(horizontal) as archive:
        assert b'<w:textDirection w:val="tbRl"' not in archive.read(
            "word/document.xml"
        )


def test_convert_docx_accepts_b6_on_a5_with_bottom_right_margins(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "b6.docx"
    _make_source_docx(source)

    result = convert_docx(
        source,
        output,
        LayoutSettings(
            imposition_mode="b6_on_a5",
            writing_mode="taiwan_vertical",
            binding_direction="right",
            page_numbers=False,
        ),
    )

    section = Document(output).sections[0]
    assert result.imposition_mode == "b6_on_a5"
    assert abs(section.page_width.cm - 14.8) < 0.02
    assert abs(section.page_height.cm - 21.0) < 0.02
    assert abs(section.left_margin.cm - 2.0) < 0.02
    assert abs(section.right_margin.cm - 0.0) < 0.02
    assert abs(section.top_margin.cm - 2.8) < 0.02
    assert abs(section.bottom_margin.cm - 0.0) < 0.02


def test_right_bound_docx_page_fields_use_mirrored_odd_even_footers(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.docx"
    output = tmp_path / "right-bound.docx"
    _make_source_docx(source)

    convert_docx(
        source,
        output,
        LayoutSettings(
            imposition_mode="single_a5",
            binding_direction="right",
            page_numbers=True,
        ),
    )

    section = Document(output).sections[0]
    assert section.footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert section.even_page_footer.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert "PAGE" in section.footer._element.xml
    assert "PAGE" in section.even_page_footer._element.xml
