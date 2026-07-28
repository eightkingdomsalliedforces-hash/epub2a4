from __future__ import annotations

from dataclasses import replace
from pathlib import Path
from typing import Iterable
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm

from .crop_frame import CropFrameLine, build_crop_frame
from .geometry import RectMm, calculate_layout
from .models import CoverElement, CoverProject, ElementKind
from .ooxml import (
    CropPercent,
    add_anchored_picture,
    add_line_shape,
    add_text_box,
    make_text_box_shape,
    mm_to_twips,
)
from .pdf_export import CoverExportError, ExportResult
from .print_plan import PrintMark, PrintPage, build_print_plan, visible_print_marks
from .project_io import validate_project


def _intersection(first: RectMm, second: RectMm) -> RectMm | None:
    left = max(first.x_mm, second.x_mm)
    top = max(first.y_mm, second.y_mm)
    right = min(first.right_mm, second.right_mm)
    bottom = min(first.bottom_mm, second.bottom_mm)
    if right <= left or bottom <= top:
        return None
    return RectMm(left, top, right - left, bottom - top)


def _element_rect(element: CoverElement) -> RectMm:
    transform = element.transform
    return RectMm(
        float(transform.x_mm),
        float(transform.y_mm),
        float(transform.width_mm),
        float(transform.height_mm),
    )


def _to_page_rect(rect: RectMm, page: PrintPage) -> RectMm:
    return RectMm(
        page.destination_rect.x_mm + rect.x_mm - page.source_rect.x_mm,
        page.destination_rect.y_mm + rect.y_mm - page.source_rect.y_mm,
        rect.width_mm,
        rect.height_mm,
    )


def _crop_for_intersection(full: RectMm, visible: RectMm) -> CropPercent:
    left = round((visible.x_mm - full.x_mm) / full.width_mm * 100000)
    top = round((visible.y_mm - full.y_mm) / full.height_mm * 100000)
    right = round((full.right_mm - visible.right_mm) / full.width_mm * 100000)
    bottom = round((full.bottom_mm - visible.bottom_mm) / full.height_mm * 100000)
    return CropPercent(
        max(0, min(99999, left)),
        max(0, min(99999, top)),
        max(0, min(99999, right)),
        max(0, min(99999, bottom)),
    )


def _configure_section(section, page: PrintPage) -> None:
    width_mm, height_mm = page.paper_size_mm
    section.orientation = (
        WD_ORIENT.LANDSCAPE if page.orientation == "landscape" else WD_ORIENT.PORTRAIT
    )
    section.page_width = mm_to_twips(width_mm)
    section.page_height = mm_to_twips(height_mm)
    safe_margin = Mm(12.7)
    section.top_margin = safe_margin
    section.right_margin = safe_margin
    section.bottom_margin = safe_margin
    section.left_margin = safe_margin
    section.header_distance = 0
    section.footer_distance = 0
    section.gutter = 0

    properties = section._sectPr
    page_size = properties.find(qn("w:pgSz"))
    if page_size is None:
        page_size = OxmlElement("w:pgSz")
        properties.insert(0, page_size)
    page_size.set(qn("w:w"), str(mm_to_twips(width_mm)))
    page_size.set(qn("w:h"), str(mm_to_twips(height_mm)))
    if page.orientation == "landscape":
        page_size.set(qn("w:orient"), "landscape")
    else:
        page_size.attrib.pop(qn("w:orient"), None)

    page_margins = properties.find(qn("w:pgMar"))
    if page_margins is None:
        page_margins = OxmlElement("w:pgMar")
        properties.append(page_margins)
    safe_margin_twips = str(mm_to_twips(12.7))
    for name in ("top", "right", "bottom", "left"):
        page_margins.set(qn(f"w:{name}"), safe_margin_twips)
    for name in ("header", "footer", "gutter"):
        page_margins.set(qn(f"w:{name}"), "0")


def _remove_default_paragraph(document: Document) -> None:
    body = document._body._element
    paragraphs = body.findall(qn("w:p"))
    if len(paragraphs) == 1 and not paragraphs[0].xpath(".//w:t"):
        body.remove(paragraphs[0])


def _add_text_element(paragraph, element: CoverElement, visible: RectMm, page: PrintPage) -> None:
    page_rect = _to_page_rect(visible, page)
    if visible != _element_rect(element):
        element = replace(
            element,
            transform=replace(
                element.transform,
                x_mm=visible.x_mm,
                y_mm=visible.y_mm,
                width_mm=visible.width_mm,
                height_mm=visible.height_mm,
            ),
        )
    add_text_box(paragraph, element, page_rect)


def _add_shape_element(paragraph, element: CoverElement, visible: RectMm, page: PrintPage) -> None:
    page_rect = _to_page_rect(visible, page)
    text = str(element.content.get("text", ""))
    fill = element.content.get("fill")
    stroke = element.content.get("stroke")
    shape = make_text_box_shape(
        shape_id=f"shape-{element.id}",
        rect=page_rect,
        rotation_deg=element.transform.rotation_deg,
        text=text,
        font_family=str(element.content.get("font_family", "sans-serif")),
        font_size_pt=float(element.content.get("font_size_pt", 8.0)),
        color=str(element.content.get("color", "#000000")),
        align=str(element.content.get("align", "center")),
        line_spacing=float(element.content.get("line_spacing", 1.0)),
        fill=str(fill) if fill else None,
        stroke=str(stroke) if stroke else None,
        z_index=element.z_index,
    )
    paragraph._p.append(shape)


def _add_elements(paragraph, project: CoverProject, page: PrintPage, drawing_id: int) -> int:
    indexed = sorted(enumerate(project.elements), key=lambda item: (item[1].z_index, item[0]))
    for _index, element in indexed:
        if element.kind is ElementKind.GUIDE and not bool(element.content.get("printable", False)):
            continue
        rect = _element_rect(element)
        visible = _intersection(rect, page.source_rect)
        if visible is None:
            continue
        if element.kind is ElementKind.IMAGE:
            path_value = element.content.get("path")
            if not isinstance(path_value, str) or not Path(path_value).is_file():
                raise CoverExportError(f"元素 {element.id} 的圖片不存在：{path_value}")
            add_anchored_picture(
                paragraph,
                Path(path_value),
                _to_page_rect(visible, page),
                _crop_for_intersection(rect, visible),
                drawing_id,
                behind_text=element.z_index < 0,
                name=element.id,
            )
            drawing_id += 1
        elif element.kind is ElementKind.TEXT:
            _add_text_element(paragraph, element, visible, page)
        elif element.kind in {
            ElementKind.SHAPE,
            ElementKind.BARCODE_PLACEHOLDER,
            ElementKind.GUIDE,
        }:
            _add_shape_element(paragraph, element, visible, page)
    return drawing_id


def _label_rect(mark: PrintMark) -> RectMm:
    width = 90.0 if mark.role == "instruction" else 60.0
    height = 7.0 if mark.role == "label" else 6.0
    return RectMm(
        mark.x1_mm - width / 2.0,
        max(0.0, mark.y1_mm - height / 2.0),
        width,
        height,
    )


def _add_print_marks(paragraph, project: CoverProject, page: PrintPage) -> None:
    for index, mark in enumerate(visible_print_marks(project, page), start=1):
        if mark.kind == "line" and mark.x2_mm is not None and mark.y2_mm is not None:
            add_line_shape(
                paragraph,
                shape_id=f"mark-{page.name}-{index}",
                x1_mm=mark.x1_mm,
                y1_mm=mark.y1_mm,
                x2_mm=mark.x2_mm,
                y2_mm=mark.y2_mm,
                behind_text=True,
                dash_style=mark.line_style,
            )
        elif mark.kind == "label":
            font_size = 10.0 if mark.role == "label" else 8.0
            shape = make_text_box_shape(
                shape_id=f"label-{page.name}-{index}",
                rect=_label_rect(mark),
                rotation_deg=0.0,
                text=mark.label,
                font_family="sans-serif",
                font_size_pt=font_size,
                color="#000000",
                align="center",
                line_spacing=1.0,
                behind_text=True,
                z_index=-1,
            )
            paragraph._p.append(shape)


def _clip_crop_line(
    line: CropFrameLine,
    source: RectMm,
) -> tuple[float, float, float, float] | None:
    if line.y1_mm == line.y2_mm:
        if not source.y_mm <= line.y1_mm <= source.bottom_mm:
            return None
        left = max(min(line.x1_mm, line.x2_mm), source.x_mm)
        right = min(max(line.x1_mm, line.x2_mm), source.right_mm)
        if right < left:
            return None
        return left, line.y1_mm, right, line.y2_mm
    if line.x1_mm == line.x2_mm:
        if not source.x_mm <= line.x1_mm <= source.right_mm:
            return None
        top = max(min(line.y1_mm, line.y2_mm), source.y_mm)
        bottom = min(max(line.y1_mm, line.y2_mm), source.bottom_mm)
        if bottom < top:
            return None
        return line.x1_mm, top, line.x2_mm, bottom
    return None


def _add_crop_frame(
    paragraph,
    project: CoverProject,
    layout,
    page: PrintPage,
) -> None:
    for index, line in enumerate(build_crop_frame(project, layout), start=1):
        clipped = _clip_crop_line(line, page.source_rect)
        if clipped is None:
            continue
        x1, y1, x2, y2 = clipped
        add_line_shape(
            paragraph,
            shape_id=f"crop-frame-{page.name}-{index}",
            x1_mm=page.destination_rect.x_mm + x1 - page.source_rect.x_mm,
            y1_mm=page.destination_rect.y_mm + y1 - page.source_rect.y_mm,
            x2_mm=page.destination_rect.x_mm + x2 - page.source_rect.x_mm,
            y2_mm=page.destination_rect.y_mm + y2 - page.source_rect.y_mm,
            behind_text=False,
            width_pt=line.width_pt,
        )


def _validate_docx(output: Path, expected_sections: int) -> None:
    try:
        reopened = Document(output)
        if len(reopened.sections) != expected_sections:
            raise CoverExportError(
                f"DOCX section 數量不一致：{len(reopened.sections)} != {expected_sections}"
            )
        with ZipFile(output) as package:
            names = set(package.namelist())
            required = {"[Content_Types].xml", "word/document.xml", "word/_rels/document.xml.rels"}
            missing = required - names
            if missing:
                raise CoverExportError("DOCX 缺少必要部件：" + "、".join(sorted(missing)))
    except (BadZipFile, KeyError, OSError) as exc:
        raise CoverExportError(f"DOCX 套件驗證失敗：{exc}") from exc


def _project_warnings(project: CoverProject) -> tuple[str, ...]:
    raw = project.background.get("warnings", ())
    if isinstance(raw, (list, tuple)):
        return tuple(str(item) for item in raw)
    return (str(raw),) if raw else ()


def export_docx(project: CoverProject, output_path: Path | str) -> ExportResult:
    """Export one editable A4 Word section for every page in the print plan."""

    validate_project(project)
    output = Path(output_path)
    if output.suffix.lower() != ".docx":
        raise ValueError("DOCX 輸出路徑必須使用 .docx 副檔名。")
    output.parent.mkdir(parents=True, exist_ok=True)
    layout = calculate_layout(project)
    plan = build_print_plan(layout)

    document = Document()
    _remove_default_paragraph(document)
    drawing_id = 1
    for index, page in enumerate(plan.pages):
        if index == 0:
            section = document.sections[0]
        else:
            section = document.add_section(WD_SECTION.NEW_PAGE)
        _configure_section(section, page)
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1
        _add_print_marks(paragraph, project, page)
        drawing_id = _add_elements(paragraph, project, page, drawing_id)
        _add_crop_frame(paragraph, project, layout, page)

    temporary = output.with_suffix(output.suffix + ".tmp")
    try:
        document.save(temporary)
        _validate_docx(temporary, len(plan.pages))
        temporary.replace(output)
    except Exception as exc:
        output.unlink(missing_ok=True)
        temporary.unlink(missing_ok=True)
        if isinstance(exc, (CoverExportError, ValueError)):
            raise
        raise CoverExportError(f"DOCX 匯出失敗：{exc}") from exc

    return ExportResult(
        path=output,
        page_count=len(plan.pages),
        mode=plan.mode,
        dpi=0,
        warnings=_project_warnings(project),
    )
