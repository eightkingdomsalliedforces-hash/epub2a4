from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from epub_a4_word.crop_marks import add_guides_to_paragraph
from epub_a4_word.docx_writer import write_docx
from epub_a4_word.models import TextBlock, TextRun
from epub_a4_word.page_placement import CropGuide
from epub_a4_word.pagination import LayoutSettings, MiniPage


def _page() -> MiniPage:
    return MiniPage(
        [TextBlock((TextRun("正文"),), style="body")],
        used_points=20.0,
        logical_page_number=1,
    )


def _header_xml(path: Path) -> str:
    with ZipFile(path) as archive:
        return "".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.startswith("word/header") and name.endswith(".xml")
        )


def _document_xml(path: Path) -> str:
    with ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


def test_page_local_drawingml_guides_keep_distinct_coordinates(
    tmp_path: Path,
) -> None:
    document = Document()
    odd = document.add_paragraph()
    even = document.add_paragraph()
    add_guides_to_paragraph(
        odd,
        (CropGuide(20.0, 0.0, 20.0, 210.0),),
        paper_width_mm=148.0,
        paper_height_mm=210.0,
        render_mode="drawingml",
        identifier_start=1,
    )
    add_guides_to_paragraph(
        even,
        (CropGuide(128.0, 0.0, 128.0, 210.0),),
        paper_width_mm=148.0,
        paper_height_mm=210.0,
        render_mode="drawingml",
        identifier_start=11,
    )
    output = tmp_path / "page-local-guides.docx"
    document.save(output)

    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")

    assert xml.count("<w:drawing") == 2
    assert f"<wp:posOffset>{20 * 36000}</wp:posOffset>" in xml
    assert f"<wp:posOffset>{128 * 36000}</wp:posOffset>" in xml
    assert 'name="epub2a4-crop-guide-1"' in xml
    assert 'name="epub2a4-crop-guide-11"' in xml


def test_b6_drawingml_guides_use_same_full_page_coordinates(tmp_path: Path) -> None:
    output = tmp_path / "drawingml-b6.docx"
    write_docx(
        [_page()],
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(
            imposition_mode="b6_on_a5",
            output_mark_mode="crop_marks",
            cut_guides=True,
            guide_render_mode="drawingml",
        ),
        imposition_mode="b6_on_a5",
    )

    xml = _document_xml(output)
    assert "<w:drawing" in xml
    assert "<v:line" not in xml
    assert xml.count('name="epub2a4-crop-guide-') == 2
    positions = re.findall(
        r'<wp:positionH relativeFrom="page"><wp:posOffset>(\d+)</wp:posOffset></wp:positionH>'
        r'.*?<wp:positionV relativeFrom="page"><wp:posOffset>(\d+)</wp:posOffset></wp:positionV>'
        r'.*?<wp:extent cx="(\d+)" cy="(\d+)"',
        xml,
        flags=re.DOTALL,
    )
    emu_per_mm = 36000
    assert positions == [
        ("0", str(28 * emu_per_mm), str(148 * emu_per_mm), "1"),
        (str(20 * emu_per_mm), "0", "1", str(210 * emu_per_mm)),
    ]
    assert _header_xml(output) == ""


def test_signature_drawingml_fold_guides_are_dashed(tmp_path: Path) -> None:
    output = tmp_path / "drawingml-signature.docx"
    write_docx(
        [_page()],
        output,
        resources={},
        media_types={},
        settings=LayoutSettings(
            imposition_mode="signature16",
            cut_guides=True,
            guide_render_mode="drawingml",
        ),
        imposition_mode="signature16",
    )

    xml = _header_xml(output)
    assert xml.count('name="epub2a4-fold-guide-') == 2
    assert xml.count('<a:prstDash val="dash"') == 2
