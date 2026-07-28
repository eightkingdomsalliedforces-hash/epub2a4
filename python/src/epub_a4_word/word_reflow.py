from __future__ import annotations

from pathlib import Path
from typing import Callable
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .converter import ConversionResult
from .pagination import LayoutSettings, resolve_layout

ProgressCallback = Callable[[int, str], None]
_SUPPORTED_MODES = {"single_a5", "single_4x6", "b6_on_a5"}


def _notify(progress: ProgressCallback | None, percent: int, message: str) -> None:
    if progress is not None:
        progress(percent, message)


def _clear_story(story) -> None:
    element = story._element
    for child in list(element):
        element.remove(child)
    element.append(OxmlElement("w:p"))


def _append_page_field(paragraph, alignment) -> None:
    paragraph.alignment = alignment
    run = paragraph.add_run()
    run.font.size = Pt(8)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.extend((begin, instruction, separate, text, end))


def _remove_page_number_restart(section) -> None:
    sect_pr = section._sectPr
    for node in list(sect_pr.findall(qn("w:pgNumType"))):
        sect_pr.remove(node)


def _set_section_text_direction(section, value: str | None) -> None:
    sect_pr = section._sectPr
    node = sect_pr.find(qn("w:textDirection"))
    if value is None:
        if node is not None:
            sect_pr.remove(node)
        return
    if node is None:
        node = OxmlElement("w:textDirection")
        sect_pr.append(node)
    node.set(qn("w:val"), value)


def _configure_section(section, settings: LayoutSettings, add_page_number: bool) -> None:
    resolved = resolve_layout(settings)
    assert resolved.paper_width_cm is not None
    assert resolved.paper_height_cm is not None
    assert resolved.page_margin_left_cm is not None
    assert resolved.page_margin_right_cm is not None
    assert resolved.page_margin_top_cm is not None
    assert resolved.page_margin_bottom_cm is not None

    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(resolved.paper_width_cm)
    section.page_height = Cm(resolved.paper_height_cm)

    section.top_margin = Cm(resolved.page_margin_top_cm)
    section.left_margin = Cm(resolved.page_margin_left_cm)
    section.right_margin = Cm(resolved.page_margin_right_cm)
    # A footer needs a small reserved strip. With page numbers disabled the
    # requested margin preset is applied exactly, including true 0 mm mode.
    section.bottom_margin = Cm(
        max(resolved.page_margin_bottom_cm, 0.62)
        if add_page_number
        else resolved.page_margin_bottom_cm
    )
    section.header_distance = Cm(0.1)
    section.footer_distance = Cm(0.12)
    section.different_first_page_header_footer = False
    _remove_page_number_restart(section)
    _set_section_text_direction(
        section,
        "tbRl" if resolved.writing_mode == "taiwan_vertical" else None,
    )

    stories = (
        section.header,
        section.first_page_header,
        section.even_page_header,
        section.footer,
        section.first_page_footer,
        section.even_page_footer,
    )
    for story in stories:
        story.is_linked_to_previous = False
        _clear_story(story)

    if add_page_number:
        odd_alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
            if resolved.binding_direction == "left"
            else WD_ALIGN_PARAGRAPH.LEFT
        )
        even_alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if resolved.binding_direction == "left"
            else WD_ALIGN_PARAGRAPH.RIGHT
        )
        _append_page_field(section.footer.paragraphs[0], odd_alignment)
        _append_page_field(
            section.even_page_footer.paragraphs[0],
            even_alignment,
        )



def _fit_inline_shapes(document, max_width_emu: int, max_height_emu: int) -> None:
    for shape in document.inline_shapes:
        width = int(shape.width)
        height = int(shape.height)
        if width <= 0 or height <= 0:
            continue
        scale = min(1.0, max_width_emu / width, max_height_emu / height)
        if scale < 1.0:
            shape.width = round(width * scale)
            shape.height = round(height * scale)


def _fit_tables(document, max_width_twips: int) -> None:
    for table_element in document._element.xpath(".//w:tbl"):
        tbl_pr = table_element.tblPr

        tbl_width = tbl_pr.first_child_found_in("w:tblW")
        if tbl_width is None:
            tbl_width = OxmlElement("w:tblW")
            tbl_pr.append(tbl_width)
        tbl_width.set(qn("w:type"), "pct")
        tbl_width.set(qn("w:w"), "5000")

        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "autofit")

        indent = tbl_pr.first_child_found_in("w:tblInd")
        if indent is not None:
            indent.set(qn("w:w"), "0")
            indent.set(qn("w:type"), "dxa")

        grid_columns = table_element.xpath("./w:tblGrid/w:gridCol")
        widths = [int(column.get(qn("w:w"), "0")) for column in grid_columns]
        total = sum(widths)
        if grid_columns and total > max_width_twips and total > 0:
            scale = max_width_twips / total
            scaled = [max(1, round(width * scale)) for width in widths]
            # Correct rounding so the final total never exceeds the page.
            overflow = sum(scaled) - max_width_twips
            if overflow > 0:
                scaled[-1] = max(1, scaled[-1] - overflow)
            for column, width in zip(grid_columns, scaled):
                column.set(qn("w:w"), str(width))

            for cell_width in table_element.xpath(".//w:tcPr/w:tcW"):
                value = int(cell_width.get(qn("w:w"), "0"))
                cell_width.set(qn("w:type"), "dxa")
                cell_width.set(qn("w:w"), str(max(1, round(value * scale))))


def _fit_body_objects(document) -> None:
    content_widths = [
        int(section.page_width - section.left_margin - section.right_margin)
        for section in document.sections
    ]
    content_heights = [
        int(section.page_height - section.top_margin - section.bottom_margin)
        for section in document.sections
    ]
    if not content_widths or not content_heights:
        return
    max_width_emu = max(1, min(content_widths))
    max_height_emu = max(1, min(content_heights))
    _fit_inline_shapes(document, max_width_emu, max_height_emu)
    _fit_tables(document, max(1, round(max_width_emu / 635)))

def _document_warnings(path: Path) -> list[str]:
    warnings: list[str] = []
    try:
        with ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
    except Exception:
        return warnings
    if b"<w:txbxContent" in xml or b"<v:textbox" in xml or b"<wps:wsp" in xml:
        warnings.append("文件包含文字方塊；縮小紙張後，絕對定位內容可能需要人工微調。")
    if b"<wp:anchor" in xml:
        warnings.append("文件包含浮動圖片或圖形；縮小紙張後，其錨點位置可能需要人工微調。")
    return warnings


def convert_docx(
    input_path: Path | str,
    output_path: Path | str,
    settings: LayoutSettings | None = None,
    progress: ProgressCallback | None = None,
) -> ConversionResult:
    settings = settings or LayoutSettings(imposition_mode="single_a5")
    source = Path(input_path)
    output = Path(output_path)

    if source.suffix.lower() != ".docx" or not source.is_file():
        raise ValueError("請選擇有效的 DOCX Word 文件。")
    if settings.imposition_mode not in _SUPPORTED_MODES:
        raise ValueError(
            "DOCX 重新排版只支援 A5、B6 置於 A5 或 4×6 英吋單頁模式。"
        )
    if output.resolve() == source.resolve():
        raise ValueError("輸出路徑不可覆蓋原始 Word 文件。")

    _notify(progress, 5, "正在讀取 Word 文件…")
    document = Document(source)
    warnings = _document_warnings(source)
    if settings.writing_mode == "taiwan_vertical":
        warnings.append(
            "直排使用 Microsoft Word 原生格式；"
            "其他閱讀器或缺少 East Asia 字型時可能替代字形。"
        )
    _notify(progress, 30, "正在保留段落、表格、圖片與原始格式…")

    document.settings.odd_and_even_pages_header_footer = settings.page_numbers
    for section in document.sections:
        _configure_section(section, settings, add_page_number=settings.page_numbers)
    _fit_body_objects(document)

    # Ask Word/LibreOffice to refresh fields such as the new PAGE field when opened.
    settings_element = document.settings._element
    update_fields = settings_element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings_element.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output.parent.mkdir(parents=True, exist_ok=True)
    _notify(progress, 75, "正在寫入重新排版的 Word…")
    document.save(output)
    _notify(progress, 100, "Word 重新排版完成。")

    title = document.core_properties.title or source.stem
    author = document.core_properties.author or ""
    image_count = len(document.inline_shapes)
    return ConversionResult(
        output_path=output,
        title=title,
        author=author,
        mini_page_count=0,
        a4_page_count=0,
        image_count=image_count,
        warnings=tuple(warnings),
        imposition_mode=settings.imposition_mode,
        paper_sheet_count=0,
        signature_count=0,
        padded_mini_page_count=0,
        source_format="docx",
    )
