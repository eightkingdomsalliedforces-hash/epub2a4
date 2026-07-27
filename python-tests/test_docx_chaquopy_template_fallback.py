from __future__ import annotations

import builtins
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.parts.hdrftr import FooterPart, HeaderPart


def _restore_descriptor(cls, name: str, value) -> None:
    setattr(cls, name, value)


def test_inline_header_footer_templates_work_without_package_files(tmp_path: Path) -> None:
    from epub_a4_word.docx_compat import install_story_template_fallbacks

    install_story_template_fallbacks(force=True)

    output = tmp_path / "fallback.docx"
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].add_run("guide")
    section.footer.paragraphs[0].add_run("page")
    document.save(output)

    with ZipFile(output) as archive:
        names = set(archive.namelist())
        assert "word/header1.xml" in names
        assert "word/footer1.xml" in names
        assert b"guide" in archive.read("word/header1.xml")
        assert b"page" in archive.read("word/footer1.xml")


def test_chaquopy_literal_parent_path_failure_installs_fallback(monkeypatch, tmp_path: Path) -> None:
    """A normalized archive member may exist while parts/../templates cannot open."""

    import epub_a4_word.docx_compat as compat

    original_header = HeaderPart.__dict__["_default_header_xml"]
    original_footer = FooterPart.__dict__["_default_footer_xml"]
    real_open = builtins.open

    def chaquopy_open(path, *args, **kwargs):
        if "parts/../templates/default-" in str(path).replace("\\", "/"):
            raise FileNotFoundError(path)
        return real_open(path, *args, **kwargs)

    monkeypatch.setattr(
        compat.hdrftr,
        "__file__",
        "/data/data/app/files/chaquopy/AssetFinder/requirements/docx/parts/hdrftr.py",
    )
    monkeypatch.setattr(builtins, "open", chaquopy_open)
    monkeypatch.setattr(compat, "_INSTALLED", False)

    try:
        assert compat.install_story_template_fallbacks() is True
        output = tmp_path / "chaquopy-fallback.docx"
        document = Document()
        document.sections[0].header.paragraphs[0].add_run("guide")
        document.sections[0].footer.paragraphs[0].add_run("page")
        document.save(output)
        with ZipFile(output) as archive:
            assert "word/header1.xml" in archive.namelist()
            assert "word/footer1.xml" in archive.namelist()
    finally:
        _restore_descriptor(HeaderPart, "_default_header_xml", original_header)
        _restore_descriptor(FooterPart, "_default_footer_xml", original_footer)
        compat._INSTALLED = False
        compat.install_story_template_fallbacks()
