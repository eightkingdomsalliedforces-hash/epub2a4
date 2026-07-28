from pathlib import Path
from zipfile import ZipFile

from docx import Document

from epub_a4_word.converter import convert_epub
from epub_a4_word.pagination import LayoutSettings


FIXTURES = Path(__file__).parents[1] / "fixtures"


def test_convert_epub_produces_editable_docx_with_text_and_image(sample_epub: Path, tmp_path: Path) -> None:
    output = tmp_path / "converted.docx"
    progress: list[int] = []

    result = convert_epub(
        sample_epub,
        output,
        LayoutSettings(font_name="Arial", page_numbers=True, imposition_mode="four_up"),
        progress=lambda percent, _message: progress.append(percent),
    )

    assert result.output_path == output
    assert result.mini_page_count >= 2
    assert result.a4_page_count == (result.mini_page_count + 3) // 4
    assert progress[0] == 5
    assert progress[-1] == 100

    document = Document(output)
    all_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "第二章" in all_text
    assert "第一章" in all_text
    with ZipFile(output) as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
    assert len(media) == 1


def test_convert_epub_signature16_reports_sides_sheets_and_signatures(sample_epub: Path, tmp_path: Path) -> None:
    output = tmp_path / "signature.docx"

    result = convert_epub(
        sample_epub,
        output,
        LayoutSettings(font_name="Arial", page_numbers=True, imposition_mode="signature16"),
    )

    expected_signatures = (result.mini_page_count + 15) // 16
    assert result.imposition_mode == "signature16"
    assert result.signature_count == expected_signatures
    assert result.a4_page_count == expected_signatures * 4
    assert result.paper_sheet_count == expected_signatures * 2
    assert result.padded_mini_page_count == expected_signatures * 16

    document = Document(output)
    assert len(document.tables) == expected_signatures * 4


def test_convert_epub_single_page_modes_report_one_sheet_per_content_page(sample_epub: Path, tmp_path: Path) -> None:
    for mode in ("single_a5", "single_4x6"):
        output = tmp_path / f"{mode}.docx"
        result = convert_epub(
            sample_epub,
            output,
            LayoutSettings(font_name="Arial", page_numbers=True, imposition_mode=mode),
        )

        assert result.imposition_mode == mode
        assert result.a4_page_count == result.mini_page_count
        assert result.paper_sheet_count == result.mini_page_count
        assert result.signature_count == 0
        assert result.padded_mini_page_count == result.mini_page_count

        document = Document(output)
        assert len(document.tables) == 1
        assert len(document.tables[0].rows) == result.mini_page_count
        assert all(len(row.cells) == 1 for row in document.tables[0].rows)


def test_convert_input_dispatches_docx_to_reflow(tmp_path: Path) -> None:
    from docx import Document
    from epub_a4_word.converter import convert_input

    source = tmp_path / "source.docx"
    document = Document()
    document.add_paragraph("同一段第一句。同一段第二句。")
    document.add_paragraph("第二段。")
    document.save(source)

    output = tmp_path / "output.docx"
    result = convert_input(
        source,
        output,
        LayoutSettings(imposition_mode="single_4x6", margin_mode="safe"),
    )

    converted = Document(output)
    assert [p.text for p in converted.paragraphs] == ["同一段第一句。同一段第二句。", "第二段。"]
    assert result.source_format == "docx"
    assert result.imposition_mode == "single_4x6"


def test_convert_epub_defaults_to_body_only_but_can_preserve_covers(
    cover_epub_factory, tmp_path: Path
) -> None:
    body_only = convert_epub(
        cover_epub_factory(),
        tmp_path / "body-only.docx",
        LayoutSettings(imposition_mode="single_a5"),
    )
    with_covers = convert_epub(
        cover_epub_factory(),
        tmp_path / "with-covers.docx",
        LayoutSettings(imposition_mode="single_a5"),
        content_only=False,
    )

    assert body_only.image_count == 0
    assert with_covers.image_count == 2


def test_vertical_epub_end_to_end_is_right_bound_and_reopenable(tmp_path: Path) -> None:
    output = tmp_path / "vertical.docx"

    result = convert_epub(
        FIXTURES / "vertical_mixed.epub",
        output,
        LayoutSettings(
            imposition_mode="signature16",
            writing_mode="taiwan_vertical",
            binding_direction="right",
            page_numbers=True,
        ),
    )

    reopened = Document(output)
    text = "\n".join(
        cell.text
        for table in reopened.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "中文 English 2026" in text
    assert len(reopened.inline_shapes) == 1
    assert result.mini_page_count >= 2
    with ZipFile(output) as archive:
        xml = archive.read("word/document.xml")
        assert b'<w:textDirection w:val="tbRl"' in xml
